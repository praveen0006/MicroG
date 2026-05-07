"""
Convert markdown documentation files to a Word document.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Define the 13 documentation files in order
DOC_FILES = [
    "01_Overview.md",
    "02_Getting_Started.md",
    "03_Configuration.md",
    "04_Data_Pipeline.md",
    "05_Feature_Engineering.md",
    "06_Models.md",
    "07_Training_Pipeline.md",
    "08_Ensembling_and_Conformal.md",
    "09_API_and_Service.md",
    "10_Dashboard.md",
    "11_Deployment_and_CI.md",
    "12_Testing_and_Quality.md",
    "13_Troubleshooting_and_Examples.md",
]

def parse_markdown_content(md_text):
    """Parse markdown into structured sections."""
    sections = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Image
        img_match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if img_match:
            img_path = img_match.group(1)
            sections.append(('image', img_path))
        # Heading
        elif line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            sections.append(('heading', level, text))
        # Code block
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            lang = line[3:].strip() if len(line) > 3 else ''
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            sections.append(('code', '\n'.join(code_lines), lang))
        # Bullet list
        elif line.strip().startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            sections.append(('list', items))
            continue
        # Paragraph
        elif line.strip():
            sections.append(('paragraph', line.strip()))
        
        i += 1
    
    return sections

def add_markdown_to_docx(doc, md_file_path):
    """Add markdown file content to Word document."""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = parse_markdown_content(content)
    base_dir = Path(md_file_path).parent.parent
    
    for section in sections:
        if section[0] == 'heading':
            _, level, text = section
            style = f'Heading {min(level, 9)}'
            doc.add_paragraph(text, style=style)
        
        elif section[0] == 'image':
            _, img_rel_path = section
            # Resolve image path relative to project root or doc folder
            img_path = Path(md_file_path).parent / img_rel_path
            if not img_path.exists():
                img_path = base_dir / img_rel_path
            
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6))
                    p = doc.add_paragraph(f"Figure: {img_path.name}")
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Error adding image {img_path.name}: {e}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_rel_path}]")

        elif section[0] == 'code':
            _, code_text, lang = section
            p = doc.add_paragraph()
            p.style = 'Normal'
            run = p.add_run(code_text.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        elif section[0] == 'list':
            _, items = section
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
        
        elif section[0] == 'paragraph':
            _, text = section
            if not re.search(r'!\[.*?\]\((.*?)\)', text): # Skip if it was already handled as image
                text = text.replace('**', '')
                text = text.replace('`', '')
                text = text.replace('_', '')
                doc.add_paragraph(text)

def main():
    """Main conversion function."""
    docs_dir = Path(__file__).parent
    output_file = docs_dir / "Sales_Forecast_Pro_Documentation.docx"
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sales Forecast Pro — Complete Documentation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary page with images from README/PROJECT_DOC
    doc.add_heading('Executive Dashboard Summary', 1)
    
    # Manually add the 4 key screenshots at the beginning
    screenshots_dir = docs_dir.parent / "screenshots"
    for img_name in ["overview.png", "forecast.png", "models.png", "national.png"]:
        img_path = screenshots_dir / img_name
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6))
            p = doc.add_paragraph(f"Dashboard View: {img_name.replace('.png', '').capitalize()}")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("") # Spacer
    
    doc.add_page_break()

    # Add each markdown file
    all_files = ["PROJECT_DOCUMENTATION.md"] + DOC_FILES
    for md_file in all_files:
        md_path = docs_dir / md_file
        
        if md_path.exists():
            print(f"Adding {md_file}...")
            add_markdown_to_docx(doc, md_path)
            doc.add_page_break()
        else:
            print(f"Warning: {md_file} not found")
    
    # Save the document
    doc.save(output_file)
    print(f"\n[DONE] Word document created: {output_file}")
    print(f"  Size: {output_file.stat().st_size / 1024:.1f} KB")

if __name__ == "__main__":
    main()
